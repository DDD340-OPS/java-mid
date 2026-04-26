from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = ROOT / "docs" / "java-mid-project-design-report.docx"

PROJECT_TREE = """src
└── com
    └── midterm
        └── ledger
            ├── LedgerApp.java
            ├── entity
            │   ├── Ledger.java
            │   ├── MonthlySummary.java
            │   ├── StatisticsSummary.java
            │   ├── Transaction.java
            │   └── TransactionType.java
            ├── service
            │   ├── FilterService.java
            │   ├── LedgerService.java
            │   └── StatisticsService.java
            ├── ui
            │   └── LedgerConsoleUI.java
            └── util
                ├── DateUtil.java
                ├── IdGenerator.java
                └── InputHelper.java
"""


def set_run_font(run, size=12, bold=False, east_asia="宋体", western="Times New Roman"):
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line_indent=True):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def add_text_paragraph(document, text, size=12, bold=False, align=None, first_line_indent=True):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    set_paragraph_format(paragraph, first_line_indent=first_line_indent)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    heading_sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(run, size=heading_sizes.get(level, 12), bold=True, east_asia="黑体")
    return paragraph


def fill_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_format(paragraph, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_paragraph_format(paragraph, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_numbered(document, text):
    paragraph = document.add_paragraph(style="List Number")
    set_paragraph_format(paragraph, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, east_asia="Consolas", western="Consolas")
    return paragraph


def configure_page(document):
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def build_cover(document):
    add_text_paragraph(document, "项目设计报告", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    add_text_paragraph(document, "Java 个人记账系统（java-mid）", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    add_text_paragraph(document, "", first_line_indent=False)
    add_text_paragraph(document, "依据当前仓库源码内容整理生成", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    add_text_paragraph(document, f"生成日期：{datetime.now().strftime('%Y-%m-%d')}", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    add_text_paragraph(document, f"项目路径：{ROOT}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    add_text_paragraph(document, "", first_line_indent=False)

    table = document.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目名称", "个人记账系统"),
        ("项目类型", "Java 控制台应用"),
        ("开发语言", "Java"),
        ("代码仓库", "DDD340-OPS/java-mid"),
        ("课程/班级/姓名/学号", "待补充"),
        ("报告用途", "课程项目设计说明"),
    ]
    for row, values in zip(table.rows, rows):
        fill_cell(row.cells[0], values[0], bold=True)
        fill_cell(row.cells[1], values[1], align=WD_ALIGN_PARAGRAPH.LEFT)

    document.add_page_break()


def build_report(document):
    add_heading(document, "摘 要", level=1)
    add_text_paragraph(
        document,
        "本项目是一个基于 Java 的个人记账系统，采用命令行交互方式实现账本管理、收支记录维护与统计分析。"
        "系统围绕“账本-交易记录-统计结果”三类核心数据展开，支持创建账本、添加记录、查看记录、修改记录、删除记录、"
        "生成统计报表、按分类统计支出、按月份汇总以及按时间范围与类别筛选数据等功能。"
        "从代码结构上看，项目采用分层思想组织实体类、业务服务类、工具类和控制台交互类，具备较好的可读性、可维护性与教学演示价值。"
        "本报告在深入分析源码的基础上，对系统需求、总体架构、模块职责、数据结构、关键流程、核心算法、运行方式及后续优化方向进行了系统说明。"
    )

    add_heading(document, "1. 项目概述", level=1)
    add_text_paragraph(
        document,
        "个人日常消费与收入记录具有高频、碎片化和统计需求明显的特点。为帮助用户快速记录收支并查看阶段性财务情况，"
        "本项目设计了一个轻量级的记账工具。项目定位不是复杂的财务系统，而是适合课程训练与面向对象编程实践的小型应用。"
    )
    add_text_paragraph(
        document,
        "系统默认在程序启动后创建名为“我的账本”的账本对象，并通过菜单循环的方式接受用户操作。所有业务围绕账本中的交易记录集合展开，"
        "操作完成后即时返回结果，适合在课堂展示、课程答辩和基础功能演示中使用。"
    )

    add_heading(document, "2. 需求分析", level=1)
    add_heading(document, "2.1 功能需求", level=2)
    functional_requirements = [
        "账本管理：支持新建账本，并在已有数据时通过二次确认避免误清空内存数据。",
        "记录新增：录入日期、交易类型、分类、金额和备注等字段，形成完整交易对象。",
        "记录查询：按日期倒序展示账本中的全部交易记录，便于快速浏览最新收支。",
        "记录修改：依据记录 ID 查找目标记录，并重新录入完整信息进行覆盖更新。",
        "记录删除：依据记录 ID 删除指定记录，并通过确认步骤降低误删风险。",
        "综合报表：统计总收入、总支出与余额，并附带支出分类占比信息。",
        "分类统计：针对支出记录按分类聚合，输出各分类金额及其占总支出的比例。",
        "月度统计：按指定年月筛选数据，统计该月收支总额、余额和记录条数。",
        "条件筛选：支持按开始日期、结束日期、分类、类型等条件进行组合筛选。",
        "演示数据：在空账本中快速导入示例记录，便于课堂演示统计与筛选功能。",
    ]
    for item in functional_requirements:
        add_bullet(document, item)

    add_heading(document, "2.2 非功能需求", level=2)
    non_functional_requirements = [
        "易用性：控制台菜单清晰，所有输入均通过提示语与循环校验进行约束。",
        "正确性：对日期、金额、空字符串和类型编码进行校验，减少非法数据进入系统。",
        "可维护性：采用 entity、service、ui、util 分包，职责边界清晰。",
        "可扩展性：为后续增加数据持久化、图形界面、导出功能等预留了结构空间。",
        "教学性：项目覆盖了枚举、集合、继承、日期 API、Stream API、异常处理等 Java 基础知识点。",
    ]
    for item in non_functional_requirements:
        add_bullet(document, item)

    add_heading(document, "3. 总体设计", level=1)
    add_heading(document, "3.1 架构设计思想", level=2)
    add_text_paragraph(
        document,
        "系统采用轻量级分层结构，可以概括为“控制台交互层 + 业务服务层 + 实体层 + 工具层”。"
        "其中，控制台交互层负责菜单展示和输入输出；业务服务层负责账本维护、统计计算与条件筛选；"
        "实体层负责描述账本、交易记录和统计结果；工具层负责日期解析、输入校验和编号生成。"
    )
    add_bullet(document, "交互层（ui）：负责用户命令接收、菜单展示、结果打印和异常提示。")
    add_bullet(document, "业务层（service）：封装增删改查、统计、筛选等核心规则。")
    add_bullet(document, "实体层（entity）：封装账本、交易记录与统计结果等领域对象。")
    add_bullet(document, "工具层（util）：提供通用输入处理、日期转换、编号生成等支撑能力。")

    add_heading(document, "3.2 目录结构设计", level=2)
    add_code_block(document, PROJECT_TREE)
    add_text_paragraph(
        document,
        "上述结构体现了“按职责分包”的组织方式。相比将所有逻辑写入 main 方法，这种设计更利于阅读、测试和后续扩展。"
    )

    add_heading(document, "3.3 运行流程概述", level=2)
    run_flow = [
        "程序入口 LedgerApp 启动后创建 LedgerConsoleUI 对象并调用 run() 方法。",
        "UI 初始化 Scanner、InputHelper、LedgerService、StatisticsService 和 FilterService。",
        "系统默认创建一个账本对象，并进入循环菜单等待用户输入操作编号。",
        "UI 根据选择调用不同服务方法完成业务处理，再将结果格式化打印到控制台。",
        "当用户输入 0 时退出循环并关闭扫描器，程序结束运行。",
    ]
    for item in run_flow:
        add_numbered(document, item)

    add_heading(document, "4. 数据结构与类设计", level=1)
    add_heading(document, "4.1 核心实体说明", level=2)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["类名", "主要字段", "职责说明", "设计特点"]
    for index, header in enumerate(headers):
        fill_cell(table.rows[0].cells[index], header, bold=True)

    entity_rows = [
        (
            "Ledger",
            "name, transactions",
            "表示一个账本，维护交易记录集合并提供增删查能力。",
            "对外返回不可修改列表，避免外部直接替换集合。",
        ),
        (
            "Transaction",
            "id, date, type, category, amount, note",
            "表示一条收支记录，是系统最核心的数据对象。",
            "在构造与 setter 中进行字段合法性约束。",
        ),
        (
            "TransactionType",
            "code, label",
            "描述交易类型，当前包括收入与支出两种取值。",
            "使用枚举避免魔法数字散落在业务代码中。",
        ),
        (
            "StatisticsSummary",
            "totalIncome, totalExpense, balance",
            "封装总收入、总支出与余额等聚合统计结果。",
            "构造时自动计算余额，减少重复逻辑。",
        ),
        (
            "MonthlySummary",
            "month, recordCount",
            "表示指定月份的统计结果，继承自 StatisticsSummary。",
            "通过继承复用收支汇总能力，并扩展月份信息。",
        ),
    ]
    for entity in entity_rows:
        row = table.add_row().cells
        for idx, value in enumerate(entity):
            fill_cell(row[idx], value, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(document, "4.2 类之间的关系", level=2)
    add_text_paragraph(
        document,
        "Ledger 与 Transaction 之间是一对多关系，一个账本包含多条交易记录；"
        "StatisticsSummary 与 MonthlySummary 之间是继承关系，月度统计在普通统计的基础上增加了月份与记录数信息；"
        "LedgerConsoleUI 依赖多个 Service 与 Util 类完成完整业务流程，是系统的调用组织中心。"
    )

    add_heading(document, "5. 核心模块设计", level=1)
    add_heading(document, "5.1 LedgerService 模块", level=2)
    add_text_paragraph(
        document,
        "LedgerService 负责账本和交易记录的核心维护逻辑，包括创建账本、添加记录、修改记录、删除记录、按 ID 查找记录、"
        "按日期排序输出记录以及加载演示数据等。该模块在新增和修改操作前统一调用 validateInput() 方法，"
        "对账本对象、日期、类型、分类和金额进行校验，从而保证进入实体对象的数据满足基本业务约束。"
    )
    add_bullet(document, "新增记录时通过 IdGenerator 生成唯一编号，并创建 Transaction 对象后加入账本集合。")
    add_bullet(document, "修改记录时先按 ID 查找，再通过 setter 完整覆盖日期、类型、分类、金额和备注。")
    add_bullet(document, "删除记录时由 Ledger.removeTransactionById() 执行实际移除。")
    add_bullet(document, "记录展示采用“日期倒序 + ID 次排序”的策略，使最新数据优先显示。")

    add_heading(document, "5.2 StatisticsService 模块", level=2)
    add_text_paragraph(
        document,
        "StatisticsService 负责所有统计相关计算，是项目中最能体现集合处理与 Stream API 使用价值的模块。"
        "其主要职责包括计算总收入、总支出和余额，按分类汇总支出金额，生成月度统计对象，以及拼接综合统计报表字符串。"
    )
    add_bullet(document, "calculateSummary() 通过两次过滤分别累计收入和支出金额。")
    add_bullet(document, "calculateExpenseByCategory() 使用 groupingBy 与 summingDouble 聚合支出分类。")
    add_bullet(document, "分类结果按金额降序排列，并收集为 LinkedHashMap 以保持展示顺序。")
    add_bullet(document, "calculateMonthlySummary() 以 YearMonth 为筛选依据，构建月度统计对象。")
    add_bullet(document, "buildFullReport() 将统计值与分类占比格式化为完整文本报表。")

    add_heading(document, "5.3 FilterService 模块", level=2)
    add_text_paragraph(
        document,
        "FilterService 负责条件筛选功能。该模块以交易记录流为基础，根据开始日期、结束日期、分类和交易类型是否为空，"
        "逐步叠加过滤条件，最后按照与列表展示一致的排序规则返回结果。"
    )
    add_bullet(document, "支持任意条件为空，说明筛选是组合式而非强制式。")
    add_bullet(document, "分类匹配使用 equalsIgnoreCase()，提高了输入容错性。")
    add_bullet(document, "filterByMonth() 通过 month.atDay(1) 与 month.atEndOfMonth() 复用通用筛选逻辑。")

    add_heading(document, "5.4 输入与工具模块", level=2)
    add_text_paragraph(
        document,
        "工具模块由 InputHelper、DateUtil 和 IdGenerator 三部分构成。InputHelper 将控制台输入逻辑集中封装，"
        "通过循环读取与异常捕获反复提示用户直到输入合法；DateUtil 负责日期和月份的解析与格式化；"
        "IdGenerator 负责生成带时间戳和顺序号的交易编号。"
    )
    add_bullet(document, "InputHelper 统一封装整数、金额、字符串、日期、月份、类型和确认输入。")
    add_bullet(document, "DateUtil 使用 LocalDate、YearMonth 与 DateTimeFormatter 处理时间数据。")
    add_bullet(document, "IdGenerator 基于当前时间和 AtomicInteger 生成交易 ID，格式类似 TXyyyyMMddHHmmss001。")

    add_heading(document, "5.5 LedgerConsoleUI 模块", level=2)
    add_text_paragraph(
        document,
        "LedgerConsoleUI 是系统的人机交互核心。该类负责展示菜单、读取用户输入、协调调用服务层并将结果打印为表格或文本。"
        "从职责划分看，它本身不承担复杂计算，而是扮演业务编排者的角色，这有助于保持交互逻辑与业务逻辑分离。"
    )
    add_bullet(document, "run() 方法负责初始化默认账本并维持主菜单循环。")
    add_bullet(document, "createLedger()、addTransaction()、updateTransaction() 等方法分别对应具体菜单项。")
    add_bullet(document, "printTransactions() 将记录列表格式化为统一表头输出。")
    add_bullet(document, "对空账本、无匹配记录、异常输入等场景提供友好提示。")

    add_heading(document, "6. 核心业务流程设计", level=1)
    add_heading(document, "6.1 新增交易流程", level=2)
    add_numbered(document, "用户选择“添加记录”菜单项。")
    add_numbered(document, "InputHelper 依次读取日期、类型、分类、金额和备注。")
    add_numbered(document, "LedgerService 对输入执行合法性校验。")
    add_numbered(document, "系统调用 IdGenerator 生成交易 ID，并构造 Transaction 对象。")
    add_numbered(document, "交易对象加入当前账本集合，控制台输出新增成功提示。")

    add_heading(document, "6.2 统计报表流程", level=2)
    add_numbered(document, "用户选择“查看统计报表”。")
    add_numbered(document, "StatisticsService 遍历账本记录，分别计算总收入与总支出。")
    add_numbered(document, "系统进一步按支出分类聚合，并计算各分类占比。")
    add_numbered(document, "buildFullReport() 将统计信息拼接为文本，返回 UI 输出。")

    add_heading(document, "6.3 条件筛选流程", level=2)
    add_numbered(document, "用户可输入开始日期、结束日期、分类或交易类型，也可直接回车跳过某项。")
    add_numbered(document, "FilterService 根据非空条件逐步缩小交易记录范围。")
    add_numbered(document, "筛选结果按日期倒序展示，同时显示结果条数。")

    add_heading(document, "7. 关键设计与实现亮点", level=1)
    highlights = [
        "枚举建模：以 TransactionType 表示收入和支出，避免直接在代码中使用 1、2 之类的魔法数字。",
        "面向对象分层：UI 层不直接操作底层集合，而是通过 Service 完成业务处理，结构清晰。",
        "不可修改集合暴露：Ledger.getTransactions() 返回不可修改视图，降低外部误操作风险。",
        "统计对象抽象：使用 StatisticsSummary 与 MonthlySummary 对聚合结果建模，便于复用和扩展。",
        "现代日期 API：使用 LocalDate 与 YearMonth 替代旧式 Date，日期语义更清晰。",
        "Stream API 运用：在统计、排序、筛选和分组逻辑中充分体现了 Java 函数式集合处理能力。",
        "演示数据机制：loadDemoData() 让系统在答辩或课堂展示时能快速进入可观察状态。",
    ]
    for item in highlights:
        add_bullet(document, item)

    add_heading(document, "8. 运行环境与使用说明", level=1)
    add_text_paragraph(
        document,
        "从源码依赖看，本项目适合在 JDK 8 及以上环境中运行。程序为单机控制台应用，不依赖数据库、中间件或第三方服务。"
        "在具备 JDK 的情况下，可先编译全部 Java 源文件，再通过主类 com.midterm.ledger.LedgerApp 启动程序。"
    )
    add_code_block(
        document,
        "Windows PowerShell 示例：\n"
        "javac -encoding UTF-8 -d out (Get-ChildItem -Recurse -Filter *.java | ForEach-Object { $_.FullName })\n"
        "java -cp out com.midterm.ledger.LedgerApp"
    )
    add_text_paragraph(
        document,
        "程序启动后，建议先体验“加载演示数据”功能，再查看统计报表、分类统计、月度统计与条件筛选，能够更完整地展示项目功能。"
    )

    add_heading(document, "9. 存在问题与改进方向", level=1)
    issues = [
        "当前数据仅保存在内存中，程序退出后记录会丢失，缺少文件或数据库持久化能力。",
        "系统尚未提供单元测试，核心逻辑虽然清晰，但缺少自动化回归验证。",
        "交易分类采用自由文本输入，若用户输入近义词或错别字，可能导致统计分散。",
        "IdGenerator 的顺序号在程序重启后会重新计数，严格意义上的跨会话唯一性仍可加强。",
        "系统目前仅支持单用户控制台交互，尚未涉及权限管理、并发控制与图形界面。",
    ]
    for item in issues:
        add_bullet(document, item)

    add_heading(document, "10. 结论", level=1)
    add_text_paragraph(
        document,
        "总体来看，本项目围绕个人记账这一明确场景，完整实现了从数据录入、数据维护到统计分析的闭环流程。"
        "其源码结构体现了较好的面向对象设计意识，既满足课程项目对功能完整性的要求，也展示了 Java 基础语法、集合框架、"
        "日期处理、Stream API 和异常控制等知识点。若后续继续补充持久化、测试和可视化界面，本系统仍具备进一步演化为更完整应用的潜力。"
    )


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_page(document)
    build_cover(document)
    build_report(document)
    document.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
